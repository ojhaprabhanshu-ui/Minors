import pymupdf
from docx import Document
import re

from ats.skills import SKILLS


# =========================================================
# SKILL NORMALIZATION
# =========================================================

SKILL_ALIASES = {

    "restful api": "REST API",
    "restful apis": "REST API",
    "rest api": "REST API",
    "rest apis": "REST API",

    "postgres": "PostgreSQL",
    "postgresql": "PostgreSQL",

    "py": "Python",

    "ml": "Machine Learning",
    "machine learning": "Machine Learning",

    "ai": "Artificial Intelligence",
    "artificial intelligence": "Artificial Intelligence",

    "nlp": "Natural Language Processing",
    "natural language processing":
        "Natural Language Processing",

    "js": "JavaScript",
    "javascript": "JavaScript",

    "ts": "TypeScript",
    "typescript": "TypeScript",

    "node": "Node.js",
    "nodejs": "Node.js",
    "node.js": "Node.js",

    "reactjs": "React",
    "react.js": "React",
    "react": "React",

    "mongo": "MongoDB",
    "mongodb": "MongoDB"

}


# =========================================================
# NORMALIZE TEXT
# =========================================================

def normalize_text(text):

    if not text:
        return ""

    text = text.replace(
        "\x00",
        " "
    )

    text = text.replace(
        "\u2022",
        " "
    )

    text = re.sub(
        r"[ \t]+",
        " ",
        text
    )

    text = re.sub(
        r"\n{3,}",
        "\n\n",
        text
    )

    return text.strip()


# =========================================================
# PDF EXTRACTION
# =========================================================

def extract_text_from_pdf(file):

    try:

        pdf = pymupdf.open(
            stream=file.read(),
            filetype="pdf"
        )

        pages = []

        for page in pdf:

            page_text = page.get_text(
                "text"
            )

            if page_text:
                pages.append(
                    page_text
                )

        pdf.close()

        text = "\n".join(pages)

        return normalize_text(text)

    except Exception as e:

        raise ValueError(
            f"Could not read PDF resume: {str(e)}"
        )


# =========================================================
# DOCX EXTRACTION
# =========================================================

def extract_text_from_docx(file):

    try:

        document = Document(file)

        parts = []

        # Paragraphs
        for paragraph in document.paragraphs:

            if paragraph.text.strip():

                parts.append(
                    paragraph.text
                )

        # Tables
        for table in document.tables:

            for row in table.rows:

                row_text = []

                for cell in row.cells:

                    if cell.text.strip():

                        row_text.append(
                            cell.text
                        )

                if row_text:

                    parts.append(
                        " ".join(row_text)
                    )

        return normalize_text(
            "\n".join(parts)
        )

    except Exception as e:

        raise ValueError(
            f"Could not read DOCX resume: {str(e)}"
        )


# =========================================================
# CANONICAL SKILL NAME
# =========================================================

def normalize_skill(skill):

    skill_lower = skill.lower().strip()

    return SKILL_ALIASES.get(
        skill_lower,
        skill
    )


# =========================================================
# SKILL EXTRACTION
# =========================================================

def extract_skills(text):

    text_lower = normalize_text(
        text
    ).lower()

    found = set()

    for skill in SKILLS:

        skill_lower = skill.lower()

        # Allow spaces, slash and hyphen noise
        escaped = re.escape(
            skill_lower
        )

        escaped = escaped.replace(
            r"\ ",
            r"[\s\-\/]+"
        )

        pattern = (
            r"(?<!\w)"
            + escaped
            + r"(?!\w)"
        )

        if re.search(
            pattern,
            text_lower
        ):

            found.add(
                normalize_skill(skill)
            )

    # Also check aliases directly
    for alias, canonical in SKILL_ALIASES.items():

        escaped = re.escape(
            alias
        )

        escaped = escaped.replace(
            r"\ ",
            r"[\s\-\/]+"
        )

        pattern = (
            r"(?<!\w)"
            + escaped
            + r"(?!\w)"
        )

        if re.search(
            pattern,
            text_lower
        ):

            found.add(
                canonical
            )

    return sorted(found)


# =========================================================
# EXPERIENCE EXTRACTION
# =========================================================

def extract_experience(text):

    text_lower = normalize_text(
        text
    ).lower()

    # Example:
    # 2 years
    # 2+ years
    # 2 yrs
    # 2.5 years experience

    year_patterns = [

        r"(\d+(?:\.\d+)?)\s*\+?\s*"
        r"(?:years?|yrs?)"
        r"(?:\s+of)?\s+"
        r"(?:experience|exp)",

        r"(\d+(?:\.\d+)?)\s*\+?\s*"
        r"(?:years?|yrs?)",

        r"(?:experience|exp)"
        r"\s*[:\-]?\s*"
        r"(\d+(?:\.\d+)?)\s*\+?\s*"
        r"(?:years?|yrs?)"

    ]

    for pattern in year_patterns:

        match = re.search(
            pattern,
            text_lower
        )

        if match:

            try:

                return float(
                    match.group(1)
                )

            except ValueError:

                pass

    # Example:
    # 24 months experience

    month_match = re.search(
        r"(\d+)\s*\+?\s*months?"
        r"(?:\s+of)?\s+"
        r"(?:experience|exp)?",
        text_lower
    )

    if month_match:

        try:

            months = float(
                month_match.group(1)
            )

            return round(
                months / 12,
                1
            )

        except ValueError:

            pass

    return 0


# =========================================================
# SECTION DETECTION
# =========================================================

def detect_sections(text):

    text_lower = normalize_text(
        text
    ).lower()

    sections = {

        "experience": False,
        "projects": False,
        "skills": False,
        "education": False,
        "certifications": False

    }

    keywords = {

        "experience": [

            "experience",
            "work experience",
            "professional experience",
            "employment",
            "work history"

        ],

        "projects": [

            "projects",
            "personal projects",
            "academic projects",
            "project experience"

        ],

        "skills": [

            "skills",
            "technical skills",
            "core skills",
            "technologies",
            "tech stack"

        ],

        "education": [

            "education",
            "academic background",
            "qualifications",
            "degree",
            "university",
            "college"

        ],

        "certifications": [

            "certifications",
            "certificates",
            "certification"

        ]

    }

    for section, words in keywords.items():

        for word in words:

            if word in text_lower:

                sections[section] = True

                break

    return sections


# =========================================================
# RESUME ANALYSIS
# =========================================================

def extract_resume_info(text):

    text = normalize_text(text)

    if not text:

        raise ValueError(
            "No readable text was found in the resume."
        )

    try:

        skills = extract_skills(
            text
        )

        experience_years = (
            extract_experience(text)
        )

        sections = detect_sections(
            text
        )

        return {

            "skills": skills,

            "experience_years":
                experience_years,

            "sections":
                sections

        }

    except Exception as e:

        raise ValueError(
            f"Resume analysis failed: {str(e)}"
        )


# =========================================================
# JOB DESCRIPTION ANALYSIS
# =========================================================

def analyze_job_description(
    job_description
):

    job_description = normalize_text(
        job_description
    )

    if not job_description:

        raise ValueError(
            "Job description is empty."
        )

    return {

        "required_skills":
            extract_skills(
                job_description
            ),

        "required_experience":
            extract_experience(
                job_description
            )

    }